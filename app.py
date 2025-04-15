from flask import Flask, render_template, request, send_file, redirect, url_for, session, flash
from functools import wraps
import psycopg2
import json, datetime
from docx import Document
from io import BytesIO

app = Flask(__name__)
app.secret_key = 'your-secret-key-123'  # Замените на надежный секретный ключ в продакшене

# Конфигурация пользователя (email: password)
USER_CREDENTIALS = {
    'user@example.com': '123456'
}

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('logged_in'):
            flash('Пожалуйста, войдите в систему для доступа к этой странице', 'danger')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        
        if email in USER_CREDENTIALS and USER_CREDENTIALS[email] == password:
            session['logged_in'] = True
            session['email'] = email
            flash('Вы успешно вошли в систему', 'success')
            return redirect(url_for('index'))
        else:
            flash('Неверный email или пароль', 'danger')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('logged_in', None)
    session.pop('email', None)
    flash('Вы успешно вышли из системы', 'success')
    return redirect(url_for('login'))

def get_db_connection():
    return psycopg2.connect(
        dbname='student_performance',
        user='postgres',
        password='123456',
        host='localhost',
        port='5432'
    )

@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute('SELECT teacher_id, full_name FROM teachers ORDER BY full_name')
    teachers = cur.fetchall()

    cur.execute('SELECT group_id, group_name FROM groups ORDER BY group_name')
    groups = cur.fetchall()

    cur.execute('SELECT DISTINCT point_number FROM control_points ORDER BY point_number')
    semesters = [(point[0], f"{point[0]} семестр") for point in cur.fetchall()]

    cur.execute('SELECT discipline_id, discipline_name FROM disciplines ORDER BY discipline_name')
    disciplines = cur.fetchall()

    results = summary = activity_types = None
    selected_teacher = selected_group = selected_semester = selected_discipline = ''
    show_details = False
    teacher_name = group_name = semester_name = discipline_name = ''

    if request.method == 'POST':
        selected_teacher = request.form.get('teacher_id', '')
        selected_group = request.form.get('group_id', '')
        selected_semester = request.form.get('semester', '')
        selected_discipline = request.form.get('discipline_id', '')
        show_details = 'show_details' in request.form

        if selected_teacher:
            cur.execute('SELECT full_name FROM teachers WHERE teacher_id = %s', (selected_teacher,))
            teacher_name = cur.fetchone()[0]

        if selected_group:
            cur.execute('SELECT group_name FROM groups WHERE group_id = %s', (selected_group,))
            group_name = cur.fetchone()[0]

        if selected_semester:
            semester_name = f"{selected_semester} семестр"

        if selected_discipline:
            cur.execute('SELECT discipline_name FROM disciplines WHERE discipline_id = %s', (selected_discipline,))
            discipline_name = cur.fetchone()[0]

        activity_types_query = """
            SELECT DISTINCT at.activity_name 
            FROM activity_types at
            JOIN point_activities pa ON at.activity_id = pa.activity_type
            JOIN student_scores ss ON pa.activity_id = ss.activity_id
            JOIN students s ON ss.student_id = s.student_id
            JOIN teachers t ON ss.teacher_id = t.teacher_id
            JOIN control_points cp ON pa.point_id = cp.point_id
            JOIN disciplines d ON cp.discipline_id = d.discipline_id
            WHERE 1=1
        """
        activity_params = []

        if selected_teacher:
            activity_types_query += " AND t.teacher_id = %s"
            activity_params.append(selected_teacher)
        if selected_group:
            activity_types_query += " AND s.group_id = %s"
            activity_params.append(selected_group)
        if selected_semester:
            activity_types_query += " AND cp.point_number = %s"
            activity_params.append(selected_semester)
        if selected_discipline:
            activity_types_query += " AND d.discipline_id = %s"
            activity_params.append(selected_discipline)

        activity_types_query += " ORDER BY at.activity_name"
        cur.execute(activity_types_query, tuple(activity_params))
        activity_types = [row[0] for row in cur.fetchall()]

        final_grades_query = """
            SELECT 
                s.full_name AS student_name,
                ROUND(fg.total_score)::integer AS final_grade,
                fg.grade AS grade_category
            FROM final_grades fg
            JOIN students s ON fg.student_id = s.student_id
            JOIN teachers t ON fg.teacher_id = t.teacher_id
            JOIN disciplines d ON fg.discipline_id = d.discipline_id
            WHERE 1=1
        """
        final_params = []

        if selected_teacher:
            final_grades_query += """
                AND s.student_id IN (
                    SELECT DISTINCT ss.student_id
                    FROM student_scores ss
                    WHERE ss.teacher_id = %s
                )
            """
            final_params.append(selected_teacher)

        if selected_group:
            final_grades_query += " AND s.group_id = %s"
            final_params.append(selected_group)
        if selected_discipline:
            final_grades_query += " AND d.discipline_id = %s"
            final_params.append(selected_discipline)

        summary_query = f"""
            WITH final_data AS (
                {final_grades_query}
            ),
            grades_grouped AS (
                SELECT 
                    CASE
                        WHEN final_grade >= 85 THEN 'Отл'
                        WHEN final_grade >= 75 THEN 'Хор'
                        WHEN final_grade >= 60 THEN 'Удовл'
                        ELSE 'Неудовл'
                    END as grade_category
                FROM final_data
            )
            SELECT 
                grade_category,
                COUNT(*) as student_count
            FROM grades_grouped
            GROUP BY grade_category
            ORDER BY 
                CASE 
                    WHEN grade_category = 'Отл' THEN 1 
                    WHEN grade_category = 'Хор' THEN 2 
                    WHEN grade_category = 'Удовл' THEN 3 
                    ELSE 4 
                END
        """
        cur.execute(summary_query, tuple(final_params))
        summary = cur.fetchall()

        if show_details:
            details_query = f"""
                SELECT 
                    student_name,
                    final_grade,
                    CASE
                        WHEN final_grade >= 85 THEN 'Отл'
                        WHEN final_grade >= 75 THEN 'Хор'
                        WHEN final_grade >= 60 THEN 'Удовл'
                        ELSE 'Неудовл'
                    END as grade_category
                FROM ({final_grades_query}) AS final_data
                ORDER BY 
                    CAST(SUBSTRING(student_name FROM 'Студент (\\d+)') AS INTEGER)
            """
            cur.execute(details_query, tuple(final_params))
            results = cur.fetchall()

    cur.close()
    conn.close()

    return render_template(
        'index.html',
        teachers=teachers,
        groups=groups,
        semesters=semesters,
        disciplines=disciplines,
        results=results,
        summary=summary,
        activity_types=activity_types,
        selected_teacher=selected_teacher,
        selected_group=selected_group,
        selected_semester=selected_semester,
        selected_discipline=selected_discipline,
        teacher_name=teacher_name,
        group_name=group_name,
        semester_name=semester_name,
        discipline_name=discipline_name,
        show_details=show_details
    )

@app.route('/export', methods=['POST'])
@login_required
def export():
    from collections import defaultdict

    teacher_name = request.form.get('teacher_name')
    group_name = request.form.get('group_name')
    semester_name = request.form.get('semester_name')
    discipline_name = request.form.get('discipline_name')

    doc = Document()
    doc.add_heading('Отчёт по успеваемости студентов', 0)

    if any([teacher_name, group_name, semester_name, discipline_name]):
        doc.add_heading('Фильтры:', level=1)
        if teacher_name:
            doc.add_paragraph(f'Преподаватель: {teacher_name}')
        if group_name:
            doc.add_paragraph(f'Группа: {group_name}')
        if semester_name:
            doc.add_paragraph(f'Семестр: {semester_name}')
        if discipline_name:
            doc.add_paragraph(f'Дисциплина: {discipline_name}')

    try:
        conn = get_db_connection()
        cur = conn.cursor()

        teacher_id = group_id = discipline_id = semester_num = None

        if teacher_name:
            cur.execute("SELECT teacher_id FROM teachers WHERE full_name = %s", (teacher_name,))
            result = cur.fetchone()
            if result:
                teacher_id = result[0]

        if group_name:
            cur.execute("SELECT group_id FROM groups WHERE group_name = %s", (group_name,))
            result = cur.fetchone()
            if result:
                group_id = result[0]

        if discipline_name:
            cur.execute("SELECT discipline_id FROM disciplines WHERE discipline_name = %s", (discipline_name,))
            result = cur.fetchone()
            if result:
                discipline_id = result[0]

        if semester_name:
            semester_num = int(semester_name.split()[0])

        query = """
            SELECT 
                s.full_name AS student_name,
                ROUND(fg.total_score)::integer AS final_score,
                CASE
                    WHEN ROUND(fg.total_score)::integer >= 85 THEN 'Отл'
                    WHEN ROUND(fg.total_score)::integer >= 75 THEN 'Хор'
                    WHEN ROUND(fg.total_score)::integer >= 60 THEN 'Удовл'
                    ELSE 'Неудовл'
                END AS grade_label
            FROM final_grades fg
            JOIN students s ON fg.student_id = s.student_id
            JOIN teachers t ON fg.teacher_id = t.teacher_id
            JOIN disciplines d ON fg.discipline_id = d.discipline_id
            JOIN control_points cp ON d.discipline_id = cp.discipline_id
            WHERE 1=1
        """
        params = []

        if teacher_id:
            query += " AND t.teacher_id = %s"
            params.append(teacher_id)
        if group_id:
            query += " AND s.group_id = %s"
            params.append(group_id)
        if discipline_id:
            query += " AND d.discipline_id = %s"
            params.append(discipline_id)
        if semester_num:
            query += " AND cp.point_number = %s"
            params.append(semester_num)

        query += " GROUP BY s.full_name, fg.total_score"

        cur.execute(query, tuple(params))
        grades_data = cur.fetchall()

        # Сводка: считаем количество по каждой категории
        grade_order = ['Отл', 'Хор', 'Удовл', 'Неудовл']
        grade_counter = defaultdict(int)
        students_with_neudovl = []

        for student_name, score, grade in grades_data:
            grade_counter[grade] += 1
            if grade == 'Неудовл':
                students_with_neudovl.append((student_name, grade))

        # Добавляем сводку по оценкам
        doc.add_heading('Сводка по оценкам', level=1)
        grades_table = doc.add_table(rows=1, cols=2)
        grades_table.style = 'Table Grid'
        hdr_cells = grades_table.rows[0].cells
        hdr_cells[0].text = 'Оценка'
        hdr_cells[1].text = 'Количество студентов'

        for grade in grade_order:
            count = grade_counter.get(grade, 0)
            if count > 0:
                row_cells = grades_table.add_row().cells
                row_cells[0].text = grade
                row_cells[1].text = str(count)

        # Таблица с Неудовл
        doc.add_heading('Студенты с неудовлетворительными оценками', level=1)
        if students_with_neudovl:
            failures_table = doc.add_table(rows=1, cols=2)
            failures_table.style = 'Table Grid'
            hdr_cells = failures_table.rows[0].cells
            hdr_cells[0].text = 'ФИО Студента'
            hdr_cells[1].text = 'Итоговая оценка'

            for student_name, final_grade in students_with_neudovl:
                row_cells = failures_table.add_row().cells
                row_cells[0].text = student_name
                row_cells[1].text = final_grade
        else:
            doc.add_paragraph('Нет студентов с неудовлетворительными оценками.')

    except Exception as e:
        print(f"Ошибка при экспорте: {e}")
        doc.add_paragraph('Ошибка при формировании отчета.')

    finally:
        if 'cur' in locals():
            cur.close()
        if 'conn' in locals():
            conn.close()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f'Отчёт_по_оценкам_{datetime.datetime.now().strftime("%Y-%m-%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(debug=True)